# document_generator.py
"""
Bu modul AI tomonidan generatsiya qilingan matnlarni Germaniya standartiga
mos Word (.docx) hujjatlariga joylaydi:
  - Lebenslauf (Europass uslubidagi jadval-based CV)
  - Motivationsschreiben (rasmiy xat formati)

Shuningdek, har bir .docx hujjatni LibreOffice (headless rejimda) orqali
.pdf formatiga konvertatsiya qilish imkonini beradi (convert_to_pdf funksiyasi).
"""

import os
import subprocess
import logging
from datetime import date
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

ACCENT_COLOR = RGBColor(0x1F, 0x3A, 0x5F)  # eski standart, orqaga moslik uchun saqlanган

# Sidebar dizayni uchun 3 rang variant: sariq, ko'k, yashil.
# Har biri: chap panel foni (hex), aksent rang (sarlavhalar/ism uchun, RGBColor),
# sidebar ichidagi oddiy matn rangi.
TEMPLATES = {
    "yellow": {
        "sidebar_bg": "1B2A3D",                          # quyuq ko'k-qora fon
        "sidebar_text": RGBColor(0xFF, 0xFF, 0xFF),       # oq matn
        "sidebar_muted": RGBColor(0xCB, 0xD5, 0xE1),
        "accent": RGBColor(0xE8, 0xA8, 0x33),             # sariq-oltin aksent
        "main_heading": RGBColor(0x1B, 0x2A, 0x3D),
    },
    "blue": {
        "sidebar_bg": "1F2937",                           # to'q kulrang-ko'k fon
        "sidebar_text": RGBColor(0xFF, 0xFF, 0xFF),
        "sidebar_muted": RGBColor(0xCB, 0xD5, 0xE1),
        "accent": RGBColor(0x5B, 0x9B, 0xD5),             # och ko'k aksent
        "main_heading": RGBColor(0x1F, 0x29, 0x37),
    },
    "green": {
        "sidebar_bg": "1E3B2E",                           # quyuq yashil fon
        "sidebar_text": RGBColor(0xFF, 0xFF, 0xFF),
        "sidebar_muted": RGBColor(0xC8, 0xDE, 0xD0),
        "accent": RGBColor(0x6F, 0xC2, 0x93),             # och yashil aksent
        "main_heading": RGBColor(0x1E, 0x3B, 0x2E),
    },
}

DEFAULT_TEMPLATE = "yellow"

SIDEBAR_WIDTH_CM = 6.0
MAIN_WIDTH_CM = 13.0


def _get_template(template: str) -> dict:
    return TEMPLATES.get(template, TEMPLATES[DEFAULT_TEMPLATE])


def _set_cell_background(cell, hex_color: str):
    """Jadval katakchasini berilgan fon rangiga bo'yaydi (hex, masalan '1F2937')."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): hex_color,
    })
    tcPr.append(shd)


def _set_cell_margins(cell, top=200, bottom=200, left=200, right=200):
    """Jadval katakchasi ichki bo'shliqlarini (padding) sozlaydi, birliklar DXA (1/20 pt)."""
    tcPr = cell._tc.get_or_add_tcPr()
    mar = tcPr.makeelement(qn("w:tcMar"), {})
    for edge, value in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        el = mar.makeelement(qn(f"w:{edge}"), {qn("w:w"): str(value), qn("w:type"): "dxa"})
        mar.append(el)
    tcPr.append(mar)


def _remove_table_borders(table):
    """Butun jadvalning barcha chegaralarini olib tashlaydi (toza, chegarasiz ko'rinish)."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = tblPr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.makeelement(qn(f"w:{edge}"), {qn("w:val"): "nil"})
        borders.append(el)
    tblPr.append(borders)


def _sidebar_heading(cell, text: str, tpl: dict):
    """Sidebar ichida kichik, aksent rangli, katta harfli sarlavha qo'shadi (masalan 'KONTAKT')."""
    p = cell.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = tpl["accent"]
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    return p


def _sidebar_line(cell, text: str, tpl: dict, size=10, bold=False, muted=False):
    """Sidebar ichiga oddiy matn qatori qo'shadi."""
    if not text:
        return None
    p = cell.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = tpl["sidebar_muted"] if muted else tpl["sidebar_text"]
    p.paragraph_format.space_after = Pt(2)
    return p


def _main_heading(cell_or_doc, text: str, tpl: dict):
    """Asosiy (o'ng) ustun uchun sarlavha qo'shadi, ostida ingichka chiziq bilan."""
    p = cell_or_doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = tpl["main_heading"]
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)

    hex_color = "%02X%02X%02X" % (tpl["main_heading"][0], tpl["main_heading"][1], tpl["main_heading"][2])
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single",
        qn("w:sz"): "6",
        qn("w:space"): "2",
        qn("w:color"): hex_color,
    })
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def convert_to_pdf(docx_path: str) -> str | None:
    """
    .docx faylni xuddi shu papkada .pdf formatiga konvertatsiya qiladi
    (LibreOffice headless rejimda ishlatiladi).

    Muvaffaqiyatli bo'lsa yaratilgan .pdf faylning to'liq yo'lini qaytaradi,
    aks holda (LibreOffice topilmasa yoki xatolik yuz bersa) None qaytaradi —
    bu holda chaqiruvchi kod faqat .docx faylni yuborish bilan davom etishi kerak.
    """
    output_dir = os.path.dirname(docx_path) or "."
    try:
        result = subprocess.run(
            [
                "soffice",
                "--headless",
                "--convert-to", "pdf",
                "--outdir", output_dir,
                docx_path,
            ],
            capture_output=True,
            text=True,
            timeout=60,
        )
        if result.returncode != 0:
            logging.warning("LibreOffice PDF konvertatsiyasi muvaffaqiyatsiz: %s", result.stderr)
            return None

        pdf_path = os.path.splitext(docx_path)[0] + ".pdf"
        if os.path.isfile(pdf_path):
            return pdf_path

        logging.warning("PDF fayl kutilgan joyda topilmadi: %s", pdf_path)
        return None
    except FileNotFoundError:
        logging.warning("LibreOffice (soffice) topilmadi - PDF konvertatsiyasi o'tkazib yuborildi.")
        return None
    except subprocess.TimeoutExpired:
        logging.warning("LibreOffice konvertatsiyasi vaqt chegarasidan oshdi: %s", docx_path)
        return None


def _set_cell_text(cell, text, bold=False, size=10, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def _styled_heading(doc, text: str, tpl: dict):
    """Shablon sozlamalariga mos rangli sarlavha qo'shadi, kerak bo'lsa
    ostiga ingichka chiziq (heading_border) chizadi."""
    h = doc.add_heading(text, level=1)
    h.runs[0].font.color.rgb = tpl["accent"]
    if tpl.get("heading_border"):
        _add_heading_border(h, tpl["accent"])
    return h


def build_cv_docx(personal: dict, cv_sections: dict, output_path: str, template: str = None):
    """
    personal: {"full_name", "address", "phone", "email", "birth_date", "nationality", "photo_path" (optional)}
    cv_sections: ai_generator.generate_cv_sections() natijasi
    template: "yellow" | "blue" | "green" - sidebar rang varianti (TEMPLATES ga qarang)

    Ikki-ustunli, to'liq balandlikdagi sidebar dizayni: chap tomonda to'q fonli
    panel (rasm, kontakt ma'lumotlari, ko'nikmalar), o'ng tomonda oq fonli
    asosiy ustun (profil, tajriba, ta'lim, tillar).
    """
    tpl = _get_template(template)
    doc = Document()

    # Sahifa marginlarini minimal qilamiz - sidebar sahifa chetigacha yetishi uchun
    for section in doc.sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0)
        section.left_margin = Cm(0)
        section.right_margin = Cm(0)

    # Bitta qatorli, 2 ustunli "asosiy" jadval - butun CV shu jadval ichida joylashadi
    outer = doc.add_table(rows=1, cols=2)
    outer.autofit = False
    outer.columns[0].width = Cm(SIDEBAR_WIDTH_CM)
    outer.columns[1].width = Cm(MAIN_WIDTH_CM)
    _remove_table_borders(outer)

    sidebar_cell = outer.rows[0].cells[0]
    main_cell = outer.rows[0].cells[1]

    sidebar_cell.width = Cm(SIDEBAR_WIDTH_CM)
    main_cell.width = Cm(MAIN_WIDTH_CM)

    _set_cell_background(sidebar_cell, tpl["sidebar_bg"])
    _set_cell_margins(sidebar_cell, top=300, bottom=300, left=250, right=250)
    _set_cell_margins(main_cell, top=300, bottom=300, left=350, right=350)

    # Boshlang'ich bo'sh paragraflarni tozalaymiz (har ikki katakda standart
    # bo'sh paragraf bo'ladi - shuni keyinroq to'ldiramiz)
    sidebar_cell.paragraphs[0].text = ""
    main_cell.paragraphs[0].text = ""

    # =====================================================================
    # SIDEBAR (chap, to'q panel)
    # =====================================================================
    photo_path = personal.get("photo_path", "")
    has_photo = bool(photo_path) and os.path.isfile(photo_path)

    first_sidebar_p = sidebar_cell.paragraphs[0]
    if has_photo:
        run = first_sidebar_p.add_run()
        run.add_picture(photo_path, width=Cm(SIDEBAR_WIDTH_CM - 1.5))
        first_sidebar_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    name_p = sidebar_cell.add_paragraph()
    name_run = name_p.add_run(personal.get("full_name", ""))
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_run.font.color.rgb = tpl["sidebar_text"]
    name_p.paragraph_format.space_before = Pt(10)
    name_p.paragraph_format.space_after = Pt(14)
    if has_photo:
        name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _sidebar_heading(sidebar_cell, "Kontakt", tpl)
    _sidebar_line(sidebar_cell, personal.get("address", ""), tpl)
    _sidebar_line(sidebar_cell, personal.get("phone", ""), tpl)
    _sidebar_line(sidebar_cell, personal.get("email", ""), tpl)

    if personal.get("birth_date") or personal.get("nationality"):
        _sidebar_heading(sidebar_cell, "Persönliche Daten", tpl)
        if personal.get("birth_date"):
            _sidebar_line(sidebar_cell, f"Geburtsdatum: {personal['birth_date']}", tpl)
        if personal.get("nationality"):
            _sidebar_line(sidebar_cell, f"Staatsangehörigkeit: {personal['nationality']}", tpl)

    skills = cv_sections.get("skills", {})
    if skills.get("languages"):
        _sidebar_heading(sidebar_cell, "Sprachen", tpl)
        for lang_item in skills["languages"]:
            _sidebar_line(sidebar_cell, lang_item, tpl, size=10)

    if skills.get("it_skills"):
        _sidebar_heading(sidebar_cell, "IT-Kenntnisse", tpl)
        for item in skills["it_skills"]:
            _sidebar_line(sidebar_cell, item, tpl, size=10)

    if skills.get("soft_skills"):
        _sidebar_heading(sidebar_cell, "Soft Skills", tpl)
        for item in skills["soft_skills"]:
            _sidebar_line(sidebar_cell, item, tpl, size=10)

    # =====================================================================
    # MAIN (o'ng, oq ustun)
    # =====================================================================
    profil_p = main_cell.paragraphs[0]
    profil_run = profil_p.add_run("Profil")
    profil_run.bold = True
    profil_run.font.size = Pt(14)
    profil_run.font.color.rgb = tpl["main_heading"]

    summary_p = main_cell.add_paragraph()
    summary_run = summary_p.add_run(cv_sections.get("profile_summary", ""))
    summary_run.font.size = Pt(10.5)
    summary_run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    summary_p.paragraph_format.space_after = Pt(8)

    _main_heading(main_cell, "Berufserfahrung", tpl)
    for job in cv_sections.get("work_experience", []):
        title_p = main_cell.add_paragraph()
        title_run = title_p.add_run(f"{job.get('position', '')} — {job.get('company', '')}")
        title_run.bold = True
        title_run.font.size = Pt(11)
        title_run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

        period_p = main_cell.add_paragraph()
        period_run = period_p.add_run(job.get("period", ""))
        period_run.italic = True
        period_run.font.size = Pt(9.5)
        period_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
        period_p.paragraph_format.space_after = Pt(2)

        for bullet in job.get("bullets", []):
            bp = main_cell.add_paragraph(style="List Bullet")
            br = bp.add_run(bullet)
            br.font.size = Pt(10)
            br.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        main_cell.add_paragraph().paragraph_format.space_after = Pt(4)

    _main_heading(main_cell, "Ausbildung", tpl)
    for edu in cv_sections.get("education", []):
        title_p = main_cell.add_paragraph()
        title_run = title_p.add_run(f"{edu.get('degree', '')} — {edu.get('institution', '')}")
        title_run.bold = True
        title_run.font.size = Pt(11)
        title_run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

        period_p = main_cell.add_paragraph()
        period_run = period_p.add_run(edu.get("period", ""))
        period_run.italic = True
        period_run.font.size = Pt(9.5)
        period_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
        period_p.paragraph_format.space_after = Pt(2)

        if edu.get("details"):
            dp = main_cell.add_paragraph()
            dr = dp.add_run(edu["details"])
            dr.font.size = Pt(10)
            dr.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        main_cell.add_paragraph().paragraph_format.space_after = Pt(4)

    output_dir = os.path.dirname(output_path)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)
    doc.save(output_path)
    return output_path


def build_motivation_letter_docx(personal: dict, recipient: dict, letter_text: str, output_path: str):
    """
    personal: {"full_name", "address", "phone", "email"}
    recipient: {"name", "institution", "address"} - kimga yo'llanayotgani (universitet/kompaniya)
    letter_text: ai_generator.generate_motivation_letter() natijasi
    """
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Yuboruvchi ma'lumotlari (yuqori o'ng burchak - DIN 5008 uslubi)
    sender_p = doc.add_paragraph()
    sender_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sender_p.add_run(
        f"{personal.get('full_name', '')}\n"
        f"{personal.get('address', '')}\n"
        f"{personal.get('phone', '')}\n"
        f"{personal.get('email', '')}"
    )

    doc.add_paragraph()

    # Qabul qiluvchi ma'lumotlari
    recipient_p = doc.add_paragraph()
    recipient_p.add_run(
        f"{recipient.get('name', '')}\n"
        f"{recipient.get('institution', '')}\n"
        f"{recipient.get('address', '')}"
    )

    doc.add_paragraph()

    # Sana
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_p.add_run(date.today().strftime("%d.%m.%Y"))

    doc.add_paragraph()

    # Xat matni - paragraflarga bo'lib qo'shamiz
    for para in letter_text.strip().split("\n\n"):
        para = para.strip()
        if para:
            doc.add_paragraph(para)

    output_dir = os.path.dirname(output_path)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)
    doc.save(output_path)
    return output_path
